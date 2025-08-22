import io
import PyPDF2
from docx import Document
import re
import logging
from typing import Dict, Any, Optional, List
from datetime import datetime

logger = logging.getLogger(__name__)

class DocumentsAgent:
    def __init__(self):
        self.supported_formats = {
            "pdf": ["application/pdf", ".pdf"],
            "docx": ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", ".docx"],
            "doc": ["application/msword", ".doc"]
        }
        
        # Maritime document patterns
        self.maritime_patterns = {
            "vessel_info": [
                r"vessel[:\s]+([A-Za-z\s\-]+?)(?:\n|$)",
                r"m/v\s+([A-Za-z\s\-]+?)(?:\n|$)",
                r"m\.v\.\s+([A-Za-z\s\-]+?)(?:\n|$)"
            ],
            "voyage_info": [
                r"voyage[:\s]+([A-Za-z0-9\s\-]+?)(?:\n|$)",
                r"voy\s+([A-Za-z0-9\s\-]+?)(?:\n|$)"
            ],
            "dates": [
                r"(\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4})",
                r"(\d{4}-\d{2}-\d{2})",
                r"(\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4})"
            ],
            "times": [
                r"(\d{1,2}:\d{2}(?:\s*[AP]M)?)",
                r"(\d{1,2}:\d{2}:\d{2})"
            ],
            "cargo_info": [
                r"cargo[:\s]+([0-9,]+(?:\.[0-9]+)?)\s*(mt|tons?|metric\s+tons?)",
                r"quantity[:\s]+([0-9,]+(?:\.[0-9]+)?)\s*(mt|tons?|metric\s+tons?)"
            ],
            "berth_info": [
                r"berth[:\s]+([A-Za-z0-9\s\-]+?)(?:\n|$)",
                r"terminal[:\s]+([A-Za-z0-9\s\-]+?)(?:\n|$)"
            ]
        }
        
        # Document type indicators
        self.document_indicators = {
            "statement_of_facts": [
                "statement of facts", "sof", "arrival", "departure", "laytime",
                "notice of readiness", "nor", "berth", "cargo operations"
            ],
            "charter_party": [
                "charter party", "cp", "charterer", "owner", "vessel",
                "cargo", "laytime", "demurrage", "despatch"
            ],
            "bill_of_lading": [
                "bill of lading", "bl", "shipper", "consignee", "notify party",
                "description of goods", "freight", "charges"
            ],
            "port_documents": [
                "port authority", "customs", "immigration", "quarantine",
                "port clearance", "berth allocation"
            ]
        }
    
    def analyze_document(self, content: bytes, filename: str, content_type: str) -> Dict[str, Any]:
        """Analyze uploaded document and extract information"""
        try:
            # Determine file format
            file_format = self._determine_file_format(filename, content_type)
            if not file_format:
                return {"error": "Unsupported file format"}
            
            # Extract text content
            if file_format == "pdf":
                text_content = self._parse_pdf(content)
            elif file_format == "docx":
                text_content = self._parse_docx(content)
            else:
                return {"error": f"Parsing not implemented for {file_format}"}
            
            if not text_content:
                return {"error": "Could not extract text from document"}
            
            # Analyze document content
            analysis_result = self._analyze_content(text_content, filename, file_format)
            
            return analysis_result
            
        except Exception as e:
            logger.error(f"Error analyzing document: {str(e)}")
            return {"error": f"Document analysis failed: {str(e)}"}
    
    def _determine_file_format(self, filename: str, content_type: str) -> Optional[str]:
        """Determine file format from filename and content type"""
        # Check content type first
        for format_name, mime_types in self.supported_formats.items():
            if content_type in mime_types:
                return format_name
        
        # Check file extension as fallback
        for format_name, mime_types in self.supported_formats.items():
            if any(filename.lower().endswith(ext) for ext in mime_types):
                return format_name
        
        return None
    
    def _parse_pdf(self, content: bytes) -> str:
        """Parse PDF content using PyPDF2"""
        try:
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text = ""
            
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error parsing PDF: {str(e)}")
            return ""
    
    def _parse_docx(self, content: bytes) -> str:
        """Parse DOCX content using python-docx"""
        try:
            doc_file = io.BytesIO(content)
            doc = Document(doc_file)
            text = ""
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + "\n"
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error parsing DOCX: {str(e)}")
            return ""
    
    def _analyze_content(self, text_content: str, filename: str, file_format: str) -> Dict[str, Any]:
        """Analyze document content and extract relevant information"""
        try:
            # Classify document type
            doc_type = self._classify_document_type(text_content)
            
            # Extract structured data
            extracted_data = self._extract_structured_data(text_content)
            
            # Generate summary
            summary = self._generate_document_summary(text_content, doc_type, extracted_data)
            
            # Calculate confidence scores
            confidence_scores = self._calculate_confidence_scores(text_content, doc_type, extracted_data)
            
            return {
                "filename": filename,
                "file_format": file_format,
                "document_type": doc_type,
                "extracted_data": extracted_data,
                "summary": summary,
                "confidence_scores": confidence_scores,
                "analysis_timestamp": datetime.now().isoformat(),
                "content_length": len(text_content),
                "content_preview": text_content[:500] + "..." if len(text_content) > 500 else text_content
            }
            
        except Exception as e:
            logger.error(f"Error analyzing content: {str(e)}")
            return {"error": f"Content analysis failed: {str(e)}"}
    
    def _classify_document_type(self, text_content: str) -> Dict[str, Any]:
        """Classify document type based on content analysis"""
        text_lower = text_content.lower()
        
        # Calculate scores for each document type
        type_scores = {}
        for doc_type, indicators in self.document_indicators.items():
            score = 0
            for indicator in indicators:
                if indicator in text_lower:
                    score += 1
            type_scores[doc_type] = score
        
        # Find the most likely document type
        if type_scores:
            most_likely = max(type_scores, key=type_scores.get)
            confidence = type_scores[most_likely] / len(self.document_indicators[most_likely])
        else:
            most_likely = "unknown"
            confidence = 0.0
        
        return {
            "primary_type": most_likely,
            "confidence": confidence,
            "all_scores": type_scores
        }
    
    def _extract_structured_data(self, text_content: str) -> Dict[str, Any]:
        """Extract structured data using pattern matching"""
        extracted_data = {}
        
        for data_type, patterns in self.maritime_patterns.items():
            extracted_data[data_type] = []
            
            for pattern in patterns:
                matches = re.findall(pattern, text_content, re.IGNORECASE)
                for match in matches:
                    if isinstance(match, tuple):
                        # Handle groups in regex
                        extracted_data[data_type].extend([m for m in match if m.strip()])
                    else:
                        # Single match
                        if match.strip():
                            extracted_data[data_type].append(match.strip())
            
            # Remove duplicates while preserving order
            seen = set()
            unique_matches = []
            for match in extracted_data[data_type]:
                if match not in seen:
                    seen.add(match)
                    unique_matches.append(match)
            extracted_data[data_type] = unique_matches
        
        return extracted_data
    
    def _generate_document_summary(self, text_content: str, doc_type: Dict[str, Any], 
                                 extracted_data: Dict[str, Any]) -> str:
        """Generate a summary of the document content"""
        primary_type = doc_type["primary_type"]
        confidence = doc_type["confidence"]
        
        summary = f"""
**Document Analysis Summary:**

**Type**: {primary_type.replace('_', ' ').title()} (Confidence: {confidence:.1%})
**Content Length**: {len(text_content)} characters

**Key Information Extracted:**
"""
        
        # Add extracted data to summary
        for data_type, values in extracted_data.items():
            if values:
                summary += f"\n**{data_type.replace('_', ' ').title()}**:\n"
                for value in values[:5]:  # Limit to first 5 values
                    summary += f"- {value}\n"
                if len(values) > 5:
                    summary += f"- ... and {len(values) - 5} more\n"
        
        # Add document-specific insights
        if primary_type == "statement_of_facts" and confidence > 0.5:
            summary += "\n**Maritime Relevance**: High - This appears to be a Statement of Facts document containing voyage operational details."
        elif primary_type == "charter_party" and confidence > 0.5:
            summary += "\n**Maritime Relevance**: High - This appears to be a Charter Party document containing contractual terms."
        elif primary_type == "bill_of_lading" and confidence > 0.5:
            summary += "\n**Maritime Relevance**: High - This appears to be a Bill of Lading document containing cargo details."
        else:
            summary += "\n**Maritime Relevance**: Low - This document doesn't appear to be primarily maritime-related."
        
        return summary.strip()
    
    def _calculate_confidence_scores(self, text_content: str, doc_type: Dict[str, Any], 
                                   extracted_data: Dict[str, Any]) -> Dict[str, float]:
        """Calculate confidence scores for different aspects of the analysis"""
        scores = {}
        
        # Document type confidence
        scores["document_type"] = doc_type["confidence"]
        
        # Data extraction confidence
        total_patterns = sum(len(patterns) for patterns in self.maritime_patterns.values())
        total_extracted = sum(len(data) for data in extracted_data.values())
        scores["data_extraction"] = min(total_extracted / max(total_patterns, 1), 1.0)
        
        # Content quality confidence
        content_score = 0.0
        if len(text_content) > 100:
            content_score += 0.3
        if len(text_content) > 500:
            content_score += 0.3
        if len(text_content) > 1000:
            content_score += 0.4
        scores["content_quality"] = content_score
        
        # Overall confidence
        scores["overall"] = (scores["document_type"] + scores["data_extraction"] + scores["content_quality"]) / 3
        
        return scores
    
    def search_document_content(self, text_content: str, search_terms: List[str]) -> Dict[str, Any]:
        """Search for specific terms in document content"""
        try:
            search_results = {}
            text_lower = text_content.lower()
            
            for term in search_terms:
                term_lower = term.lower()
                matches = []
                
                # Find all occurrences
                start = 0
                while True:
                    pos = text_lower.find(term_lower, start)
                    if pos == -1:
                        break
                    
                    # Extract context around the match
                    context_start = max(0, pos - 50)
                    context_end = min(len(text_content), pos + len(term) + 50)
                    context = text_content[context_start:context_end]
                    
                    matches.append({
                        "position": pos,
                        "context": context,
                        "highlighted": context.replace(term, f"**{term}**")
                    })
                    
                    start = pos + 1
                
                search_results[term] = {
                    "count": len(matches),
                    "matches": matches
                }
            
            return {
                "search_terms": search_terms,
                "results": search_results,
                "total_matches": sum(result["count"] for result in search_results.values())
            }
            
        except Exception as e:
            logger.error(f"Error searching document content: {str(e)}")
            return {"error": f"Search failed: {str(e)}"}
    
    def process_query(self, query: str) -> str:
        """Process natural language document-related queries"""
        query_lower = query.lower()
        
        if "document" in query_lower or "file" in query_lower or "upload" in query_lower:
            if "format" in query_lower or "supported" in query_lower:
                return f"""
**📄 Supported Document Formats:**

**PDF Documents:**
- Content type: application/pdf
- Extensions: .pdf
- Features: Full text extraction, page analysis

**Word Documents:**
- Content type: application/vnd.openxmlformats-officedocument.wordprocessingml.document
- Extensions: .docx
- Features: Text extraction, table parsing

**Legacy Word Documents:**
- Content type: application/msword  
- Extensions: .doc
- Features: Basic text extraction

**Maritime Document Types I Can Analyze:**
- Statements of Facts (SoF)
- Charter Parties
- Bills of Lading
- Port documents
- General maritime correspondence
"""
            elif "analysis" in query_lower or "parse" in query_lower:
                return """
**📄 Document Analysis Capabilities:**

**What I Extract:**
- Vessel information and names
- Voyage details and numbers
- Dates and times
- Cargo quantities and types
- Berth and terminal information
- Port details

**Analysis Features:**
- Document type classification
- Maritime content detection
- Structured data extraction
- Confidence scoring
- Content summarization
- Search functionality

**Upload a document to get started!**
"""
            else:
                return """
**📄 Maritime Document Analyzer:**

I can parse and analyze various maritime documents to extract key information.

**Supported Formats:**
- PDF files (.pdf)
- Word documents (.docx, .doc)

**Document Types:**
- Statements of Facts (SoF)
- Charter Parties
- Bills of Lading
- Port documents
- General maritime correspondence

**Ask me about:**
- "supported formats" - File types I can handle
- "analysis capabilities" - What I can extract
- "upload document" - How to analyze files
"""
        else:
            return """
**📄 Maritime Document Processing:**

I provide comprehensive document analysis for maritime operations.

**Features:**
- Multi-format support (PDF, DOCX, DOC)
- Maritime content recognition
- Structured data extraction
- Document classification
- Content search and analysis

Upload a maritime document to get started with analysis!
""" 